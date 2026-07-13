from __future__ import annotations

import shutil
import zipfile
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Inches

from .utils import contained_path, ensure_dir, read_json, read_jsonl, validate_opaque_id, write_json
from .validate import parse_translated_units


def assemble_docx(package: Path, run_id: str, out: Path) -> dict[str, Any]:
    package = package.resolve()
    out = out.resolve()
    if out == package or package in out.parents:
        raise ValueError('assembly output must be outside the package')
    run_id = validate_opaque_id(run_id, label='translation run_id')
    run_dir = contained_path(package / 'translation_runs', run_id)
    doc = Document()
    image_blocks = {img.get('marker'): img for img in read_jsonl(package / 'ledger' / 'image_blocks.jsonl')}
    image_count = 0
    unit_count = 0
    for chapter in sorted((run_dir / 'translated_md').glob('chapter_*.md')):
        doc.add_heading(chapter.stem.replace('_', ' ').title(), level=1)
        for line in chapter.read_text(encoding='utf-8').splitlines():
            if line.startswith('[') and '] ' in line:
                unit_count += 1
                doc.add_paragraph(line.split('] ', 1)[1])
            elif line.startswith('[[IMAGE '):
                image_count += 1
                img = image_blocks.get(line.strip())
                asset = None
                if img and img.get('asset_path'):
                    asset_locator = str(img['asset_path'])
                    if asset_locator.startswith('runs/'):
                        asset = contained_path(package, asset_locator)
                    else:
                        source_tree = (package / 'source' / 'epub_tree').resolve()
                        asset = contained_path(source_tree, asset_locator)
                    if asset.is_symlink():
                        raise ValueError(f'legacy image asset cannot be a symlink: {img["asset_path"]}')
                    if img.get('asset_sha256') and asset.is_file():
                        from .utils import sha256_file
                        if sha256_file(asset) != img['asset_sha256']:
                            raise ValueError(f'legacy image asset hash mismatch: {img["asset_path"]}')
                if asset and asset.exists():
                    try:
                        doc.add_picture(str(asset), width=Inches(4.5))
                    except Exception:
                        doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph(line.strip())
    ensure_dir(out.parent)
    doc.save(out)
    audit = {
        'status': 'COMPATIBILITY_ONLY_NEEDS_REVIEW' if out.exists() else 'FAIL_REVIEW',
        'output': str(out), 'units': unit_count, 'images': image_count,
        'hard_blockers': ['legacy_reinsertion_not_publication_validated'] if out.exists() else ['docx_assembly_failure'],
        'compatibility_only': True,
    }
    write_json(package / 'audit' / 'docx_assembly_audit.json', audit)
    return audit


def _visible_text_nodes(soup: BeautifulSoup):
    body = soup.find('body') or soup
    for node in body.descendants:
        if isinstance(node, NavigableString) and str(node).strip() and getattr(node.parent, 'name', None) not in {'script', 'style', 'title'}:
            yield node


def reinsert_epub(package: Path, run_id: str, out: Path) -> dict[str, Any]:
    package = package.resolve()
    out = out.resolve()
    if out == package or package in out.parents:
        raise ValueError('reinsertion output must be outside the package')
    run_id = validate_opaque_id(run_id, label='translation run_id')
    source_tree = package / 'source' / 'epub_tree'
    if not source_tree.is_dir() and (package / 'package_manifest.json').exists():
        manifest = read_json(package / 'package_manifest.json')
        active_run = str(manifest.get('active_run_id', ''))
        active_run = validate_opaque_id(active_run, label='active run_id')
        source_tree = contained_path(package, 'runs', active_run, 'assets', 'epub_tree')
    if not source_tree.is_dir():
        raise FileNotFoundError(f'legacy EPUB source tree not found: {source_tree}')
    build = contained_path(package / 'build', f'epub_{run_id}')
    if build.exists():
        shutil.rmtree(build)
    shutil.copytree(source_tree, build)
    unit_files = sorted((package / 'translation_units').glob('chapter_*.json'))
    run_dir = contained_path(package / 'translation_runs', run_id)
    replaced = 0
    touched = set()
    for uf in unit_files:
        data = read_json(uf)
        translated_path = run_dir / 'translated_md' / f"chapter_{int(data['chapter_index']):03d}.md"
        got_units, _ = parse_translated_units(translated_path)
        translated_map = {}
        for line in translated_path.read_text(encoding='utf-8').splitlines() if translated_path.exists() else []:
            if line.startswith('[') and '] ' in line:
                uid, text = line.split('] ', 1)
                translated_map[uid.strip('[')] = text
        by_href: dict[str, list[tuple[str, str]]] = {}
        for u in data.get('units', []):
            if u.get('href') and u['unit_id'] in translated_map:
                by_href.setdefault(u['href'], []).append((u['unit_id'], translated_map[u['unit_id']]))
        for href, items in by_href.items():
            matches = list(build.rglob(Path(href).name))
            if not matches:
                continue
            path = matches[0]
            soup = BeautifulSoup(path.read_text(encoding='utf-8', errors='replace'), 'lxml-xml')
            nodes = list(_visible_text_nodes(soup))
            for node, (_, text) in zip(nodes, items):
                node.replace_with(text)
                replaced += 1
            path.write_text(str(soup), encoding='utf-8')
            touched.add(str(path.relative_to(build)))
    ensure_dir(out.parent)
    mimetype = build / 'mimetype'
    with zipfile.ZipFile(out, 'w') as zf:
        if mimetype.exists():
            zf.write(mimetype, 'mimetype', compress_type=zipfile.ZIP_STORED)
        for p in sorted(build.rglob('*')):
            if p.is_file() and p != mimetype:
                zf.write(p, str(p.relative_to(build)))
    audit = {
        'status': 'COMPATIBILITY_ONLY_NEEDS_REVIEW' if out.exists() else 'FAIL_REVIEW',
        'output': str(out), 'replaced_text_nodes': replaced, 'touched_files': sorted(touched),
        'hard_blockers': ['legacy_reinsertion_not_publication_validated'] if out.exists() else ['epub_packaging_failure'],
        'compatibility_only': True,
    }
    write_json(package / 'audit' / 'epub_reinsertion_audit.json', audit)
    return audit
